#!/usr/bin/env python3
"""
DOCX extractor

Responsibilities
---------------
1) Full-text extraction for .docx files (body paragraphs/tables and section
   headers/footers).
2) Lightweight metadata extraction when `metadata_only=True`.

Design notes
------------
- Keep a single public function: `extract_docx(path, metadata_only=False)`.
- When `metadata_only` is True, return a dict of normalized metadata.
- When False, return a single string with text content.
- Raise ExtractionError when the package cannot be opened; tolerate malformed
  individual metadata fields and document blocks.
"""

from __future__ import annotations

from typing import Dict, Any
from datetime import datetime
import os

from .errors import ExtractionError

try:
    from docx import Document  # python-docx
except Exception:
    Document = None  # graceful fallback


def _safe_iso(dt) -> str | None:
    """Convert docx datetime to readable ISO string; return None if not set."""
    try:
        if not dt:
            return None
        if isinstance(dt, datetime):
            return dt.isoformat()
        # Some builds may present as python-docx DateTime object compatible with str()
        return str(dt)
    except Exception:
        return None


def _extract_block_text(block) -> list[str]:
    """Paragraphs + table cells from anything with .paragraphs/.tables —
    used for both the document body and each section's header/footer."""
    chunks = []
    try:
        for p in block.paragraphs:
            if p.text:
                chunks.append(p.text)
    except Exception:
        pass
    try:
        for table in block.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        chunks.append(cell.text)
    except Exception:
        pass
    return chunks


def _extract_docx_text(doc) -> str:
    """Aggregate visible text from paragraphs, tables, and every section's
    header/footer (PII in a header/footer is otherwise silently dropped)."""
    chunks = _extract_block_text(doc)

    try:
        for section in doc.sections:
            chunks.extend(_extract_block_text(section.header))
            chunks.extend(_extract_block_text(section.footer))
    except Exception:
        pass

    return "\n".join(chunks).strip()


def _extract_docx_coreprops(doc) -> Dict[str, Any]:
    """Normalize a small, useful subset of core properties."""
    meta: Dict[str, Any] = {}
    try:
        cp = getattr(doc, "core_properties", None)
        if not cp:
            return meta

        # Normalize the metadata keys used by reports.
        meta_map = {
            "author": cp.author,
            "title": cp.title,
            "subject": cp.subject,
            "category": cp.category,
            "comments": cp.comments,
            "created": _safe_iso(cp.created),
            "modified": _safe_iso(cp.modified),
            "last_modified_by": cp.last_modified_by,
        }

        # Drop unset/empty values to keep reports clean
        for k, v in meta_map.items():
            if v not in (None, "", "None"):
                meta[k] = v
    except Exception:
        # Preserve properties collected before a malformed field failed.
        pass

    return meta


def extract_docx(path: str, metadata_only: bool = False):
    """
    Public API used by discovery.py
    - metadata_only=True  -> returns Dict[str, Any] with normalized metadata
    - metadata_only=False -> returns str with concatenated human text

    HARDENED VERSION - Now handles:
    - None input
    - Non-string types
    - Empty strings
    - Invalid file paths
    """
    # FIX 1: Validate input type
    if not isinstance(path, str):
        raise ExtractionError(f"TypeError: path must be str, got {type(path).__name__}")

    # FIX 2: Validate non-empty path
    if not path or not path.strip():
        raise ExtractionError("ValueError: path cannot be empty")

    # FIX 3: Check library is installed
    if Document is None:
        raise ExtractionError("RuntimeError: python-docx is not installed")

    # FIX 4: Verify file exists before opening
    if not os.path.isfile(path):
        raise ExtractionError("FileNotFoundError: file does not exist")

    try:
        doc = Document(path)
    except Exception as exc:
        raise ExtractionError.from_exception(exc) from exc

    if metadata_only:
        return _extract_docx_coreprops(doc)

    return _extract_docx_text(doc)
